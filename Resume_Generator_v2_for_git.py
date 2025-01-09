import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from docx.shared import Pt, RGBColor, Cm
from docx2pdf import convert
from datetime import date
import tempfile
from docx.enum.style import WD_STYLE_TYPE


Docx_Name = "Surya_Resume_Updated_On_" + str(date.today()) + ".docx"
Pdf_Name = "Surya_Resume_Updated_On_" + str(date.today()) + ".pdf"
#print(Pdf_Name)


#Perosnal Details

Name = "XXXXX"
phone_symbol = "\U0001F4DE"
Email_symbol = "\U0001F4E7"
url_symbol = "\U0001F517"
Contact_Number = f"{phone_symbol} (+xx) xxxxx xxxxx"
Email_ID = f"xxxxx@gmail.com"
LinkedIn_ID = f"www.linkedin.com/"

def document_colour(doc):
    background = OxmlElement('w:background')
    background.set(qn('w:color'), 'F3F9FB')     # F0F8FA Define black background
    doc.element.insert(0, background)      # Insert it
    background_shp = OxmlElement('w:displayBackgroundShape') # Setting to use my background
    doc.settings.element.insert(0, background_shp)      # Apply setting
                
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    #new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.name = 'Calibri'
    new_run.font.size = Pt(10)
    new_run.font.color.rgb = RGBColor(0, 0, 255)
    new_run.font.underline = False

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def Level1_customization(para):
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = para.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(13)
    para.bold = True
    para.font.color.rgb = RGBColor(97, 182, 205) #RGBColor(88, 50, 164) 


def Custom_Style(para):
    # Add an indent to the bullet
    p_format = para.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
    
    para = para.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(11)
    para.font.color.rgb = RGBColor(0x00, 0x00, 0x00) 

def font_customization(para):
    #para = para.runs[0]
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para = para.runs[0] if para.runs else paragraph.add_run()
    para.font.name = 'Calibri'
    para.font.size = Pt(10)
    para.bold = False
    para.font.color.rgb = RGBColor(0, 0, 0) 
    
    
        
def create_resume():
    # Create a new Word Document
    doc = Document()
    
    #Set document layout, background colours, etc
    document_colour(doc)
    
    # Define a custom style for headings
    heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Calibri'
    heading_style.font.size = Pt(16)
    heading_style.font.color.rgb = RGBColor(0, 0, 0) 
    heading_style.font.bold = True
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_bullet_point(doc, bold_text, regular_text, font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
        para = doc.add_paragraph('', style='List Bullet')
        
        # Add an indent to the bullet
        p_format = para.paragraph_format
        p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
        p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
        
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    
    def add_para(doc, bold_text, regular_text, font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
        para = doc.add_paragraph('')
        
        # Add an indent to the bullet
        p_format = para.paragraph_format
        #p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
        p_format.first_line_indent = Cm(1.0)  # Negative indent for hanging bullet
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
    
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    def add_company_details(doc, bold_text, regular_text, bold_text_2,font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
        para = doc.add_paragraph('')
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
        
        bold_run = para.add_run(bold_text_2)
        bold_run.bold = True
        bold_run.italics = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
    
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = 'Professional Resume of Surya Prakash Murugan'
    core_properties.author = 'Surya Prakash Murugan'
    core_properties.subject = 'Resume'
    
    #set header & Adjust page margins (top and bottom)
    
    section = doc.sections[0]# Choosing the top most section of the page
    section.top_margin = Pt(0.7)  # Reduce the top margin for header
    section.bottom_margin = Pt(0.7)  # Reduce the bottom margin for footer
     
    # Selecting the header
    header = section.header
        
    # Add content to header
    header_para = header.paragraphs[0]
    header_para.text = ">>> This resume was generated entirely in Python. For full sourcecode "
    d = add_hyperlink(header_para,'Click Here' ,'https://github.com/Surya-analyst/Python_Scripts_Git/blob/main/Resume_Generator_for_git.py')
 
    # Add header with name and contact details
    doc.add_heading("\n" + Name, level=1).style = 'CustomHeading'
    Skills = doc.add_paragraph('Data Analyst | Excel | VBA | SQL | Python | Power BI | Tableau')
    Level1_customization(Skills)
    
    #Add contact details
    contact_paragraph = doc.add_paragraph()
    contact_paragraph.add_run(Contact_Number + " | " + Email_symbol + " ")
    add_hyperlink(contact_paragraph, Email_ID, 'mailto:suryaprakashmurugan95@gmail.com')
    contact_paragraph.add_run(" | " + url_symbol + " ")
    add_hyperlink(contact_paragraph, LinkedIn_ID, 'https://www.linkedin.com/in/msp1995/')
    font_customization(contact_paragraph)
   
    
    # Add Professional Summary
    doc.add_heading('Professional Summary', level=2)
    para = add_para(doc,'','Results-driven Data Analyst with 6 years of experience in leveraging data to drive actionable insights and strategic decisions. Proficient in Data Fetching, Data cleansing, data preparation and data visualization using tools such as SQL, Excel, VBA, Python, Tableau and PowerBI. Adept at collaborating with cross-functional teams to identify key business challenges and implement data-driven solutions that optimize performance.')
   
    # Add Core Skills
    doc.add_heading('Core Skills', level=2)
    add_bullet_point(doc, '', 'Data Analysis & Reporting (SQL, Excel, Tableau, PowerBI)')
    add_bullet_point(doc, '', 'Data Automation & Process Optimization (VBA, Python, Power Query)')
    add_bullet_point(doc, '', 'Project Coordination & Documentation')
    add_bullet_point(doc, '', 'Stakeholder Communication & Collaboration')

    # Add Skills section with bar graph
    doc.add_heading('Skill Rating', level=2)
    skills = ["Excel", "MySQL", "VBA", "Python", "Power BI", "Tableau"]
    ratings = [5, 4, 3, 3, 3, 3]
    max_rating = 5  # Maximum possible rating

    # Colors for the bars
    bar_colors = ['#009879'] * len(skills)  # Main bar color
    bg_colors = ['#b7e4c7'] * len(skills)  # Background bar color

    # Plot settings
    plt.figure(figsize=(4, 2))
    bar_height = 0.7  # Adjust height to reduce spacing (default is typically 0.8)
    for i, (skill, rating) in enumerate(zip(skills, ratings)):
        # Draw background bar
        plt.barh(y=i, width=max_rating, color=bg_colors[i], height=bar_height, edgecolor='none')
        # Draw rating bar
        plt.barh(y=i, width=rating, color=bar_colors[i], height=bar_height, edgecolor='none')
        # Add text for the rating and skill
        plt.text(rating - 0.9, i, f"{rating} / {max_rating}", va='center', ha='center', color='white', fontsize=9, weight='bold')
        plt.text(-0.99, i, skill, va='center', ha='left', color='black', fontsize=9)

    # Remove axes and labels for clean look
    plt.gca().invert_yaxis()  # Invert y-axis for descending order
    plt.axis('off')  # Hide the axes

    # Set background color
    plt.gcf().patch.set_facecolor('#F3F9FB')  # Light blue background

    # Show or Save the chart as an image
    plt.tight_layout()
    #plt.show()
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path = tmp_file.name
        plt.savefig(chart_path)
    plt.close()
    
    # Add the chart to the Word document
    doc.add_picture(chart_path, width=doc.sections[0].page_width * 0.4)  # Scale image to fit page
    
    # Add Professional Experience
    doc.add_heading('Professional Experience', level=2)
    doc.add_paragraph('Assistant Manager - MIS', style='Heading 3')
    add_company_details(doc,'Standard Chartered Bank - GBS ', 'Chennai, ','07/2023 – Present')
    add_bullet_point(doc, '','Automated daily report consolidation and data extraction using Power Query, improving data accuracy and efficiency.')
    add_bullet_point(doc, '', 'Developed solutions with Excel User Forms, VBA, and Power Query, enhancing process speed and reliability.')
    add_bullet_point(doc, '', 'Created dynamic tools for seamless data management between MS Access and Excel.')
    add_bullet_point(doc, '', 'Implemented automated reporting for trend analysis and productivity, enabling data-driven decision-making.')
  
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph('Data Analyst', style='Heading 3')
    add_company_details(doc,'Optum Health Care Business Services & Technology, ','Chennai, ','03/2021 – 07/2023')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, '', 'Extracted, cleaned, and analyzed large datasets, ensuring quality data for business insights.')
    add_bullet_point(doc, '', 'Led project planning, task scheduling, and documentation for healthcare solutions.')
    add_bullet_point(doc, '', 'Coordinated with cross-functional teams to deliver client projects successfully.')
    add_bullet_point(doc, '', 'Developed SQL queries and automated reporting to support management decisions.')
    
    doc.add_paragraph('Technical Process Specialist', style='Heading 3')
    add_company_details(doc,'Infosys, ','Bangalore, ','09/2020 – 03/2021')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, '', 'Extracted and analyzed data from CMS platforms for custom reporting.')
    add_bullet_point(doc, '', 'Built Excel dashboards and managed MySQL databases for data operations.')
    add_bullet_point(doc, '', 'Delivered stakeholder reports, ensuring clear communication of insights.')

    doc.add_paragraph('Project Coordinator', style='Heading 3')
    add_company_details(doc,'Origin Learning Solutions Pvt Ltd, ','Chennai, ','10/2018 – 01/2020')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, '', 'Managed project schedules, resources, and milestones for smooth execution.')
    add_bullet_point(doc, '', 'Collaborated with stakeholders to define objectives and monitor progress.')
    add_bullet_point(doc, '', 'Developed risk assessments and contingency plans to mitigate project issues.')
    
    doc.add_paragraph('Project Associate', style='Heading 3')
    add_company_details(doc,'Emerson Automation Solutions Pvt Ltd, ','Chennai, ','10/2018 – 01/2020')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, '', 'Oversaw PMO tasks, delivery schedules, and technical document submissions.')
    add_bullet_point(doc, '', 'Maintained client relationships by addressing feedback and meeting deadlines.')
   
    # Add Education section
    doc.add_heading('Education', level=2)
    c = doc.add_paragraph('Bachelor of Engineering', style='Heading 3')
    add_para(doc,'','R.M.D Engineering College, Chennai, 2017')
    # Add an indent to the bullet
    p_format = c.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    
    c = doc.add_paragraph('Higher Secondary School Certificate', style='Heading 3')
    add_para(doc,'','SRV Boys Higher Secondary School, Rasipuram, 2013')
    p_format = c.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)

    doc.add_heading('Additional Training or Certifications', level=2)
    c = doc.add_paragraph('MySQL for Data Analytics and Buisness Intelligence (2020), ', style='List Bullet')
    d = add_hyperlink(c,'Certificate Link' ,'https://www.udemy.com/certificate/UC-3c8b12ac-16ab-441c-a5b8-a71f287c31c1/')
    c = Custom_Style(c)

    c = doc.add_paragraph('Master Python by Coding 100 Practical Problems (2024), ', style='List Bullet')
    d = add_hyperlink(c,'Certificate Link' ,'https://www.udemy.com/certificate/UC-f4484a94-e2bb-42b9-80f8-efff2f20ae24/')
    c = Custom_Style(c)

    # Add Languages (Optional)  
    doc.add_heading('Languages', level=2)
    c = doc.add_paragraph('Tamil (Proficient)', style='List Bullet')
    d = doc.add_paragraph('English (Proficient)', style='List Bullet')
    c = Custom_Style(c)
    d = Custom_Style(d)

    # Save the document as temp file to generate pdf
    
#    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
#        filename = tmp_file.name
#        doc.save(filename)
    plt.close()
    
    #convert the docx temp file to pdf
    #convert(filename, Resume_Name)
    
    #file_name = "Professional_Resume_Template.docx"
    doc.save(Docx_Name)
    convert(Docx_Name, Pdf_Name)
    print(f"Resume template saved as {Pdf_Name}")

# Call the function to create the resume
create_resume()