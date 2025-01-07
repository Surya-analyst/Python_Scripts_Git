import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from docx.shared import Pt, RGBColor, Cm
from docx2pdf import convert
from datetime import date
import tempfile

Resume_Name = "Surya_Resume_Updated_On_" + str(date.today()) + ".pdf"
#print(Resume_Name)


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    #new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.name = 'Calibri'
    new_run.font.size = Pt(10)
    new_run.font.color.rgb = RGBColor(0, 0, 255)
    new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def Custom_Style(para):
    # Add an indent to the bullet
    p_format = para.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
    
    para = para.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(10)
    para.font.color.rgb = RGBColor(0x00, 0x00, 0x00) 
    
        
def create_resume():
    # Create a new Word Document
    doc = Document()
    
    def add_bullet_point(doc, bold_text, regular_text, font_name='Calibri', font_size=10, font_color=(0, 0, 0)):
        para = doc.add_paragraph('', style='List Bullet')
        
        # Add an indent to the bullet
        p_format = para.paragraph_format
        p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
        p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
        
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    
    def add_para(doc, bold_text, regular_text, font_name='Calibri', font_size=10, font_color=(0, 0, 0)):
        para = doc.add_paragraph('')
        
        # Add an indent to the bullet
        p_format = para.paragraph_format
        #p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
        p_format.first_line_indent = Cm(1.0)  # Negative indent for hanging bullet
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
    
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    def add_company_details(doc, bold_text, regular_text, font_name='Calibri', font_size=10, font_color=(0, 0, 0)):
        para = doc.add_paragraph('')
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
    
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    
    # Set document properties
    core_properties = doc.core_properties
    core_properties.title = 'Professional Resume of ####'
    core_properties.author = '####'
    core_properties.subject = 'Resume'
    
    #set header & Adjust page margins (top and bottom)
    
    section = doc.sections[0]# Choosing the top most section of the page
    section.top_margin = Pt(0.7)  # Reduce the top margin for header
    section.bottom_margin = Pt(0.7)  # Reduce the bottom margin for footer
     
    # Selecting the header
    header = section.header
        
    # Add content to header
    header_para = header.paragraphs[0]
    header_para.text = ">>> This resume was generated entirely in Python. For full sourcecode "
    d = add_hyperlink(header_para,'Click Here' ,'https://github.com/Surya-analyst/Python_Scripts_Git/blob/main/Resume_Generator_for_git.py')
 
    # Add header with name and contact details
    doc.add_heading('#####', level=1)
    contact_paragraph = doc.add_paragraph('Data Analyst\n')
    para = contact_paragraph.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(13)
    para.bold = True
    para.font.color.rgb = RGBColor(88, 50, 164) 
    
    add_hyperlink(contact_paragraph, '#####@gmail.com\n' ,'mailto:###@gmail.com')
    add_hyperlink(contact_paragraph,'LinkedIn Profile Link','https://www.linkedin.com/')
    run = contact_paragraph.add_run('\n(+91) ##### #####')
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(10)
    run = contact_paragraph.add_run('\nCity - Pin Code, State')
    run_font = run.font
    run_font.name = 'Calibri'
    run_font.size = Pt(10)
    
    # Add Professional Summary
    doc.add_heading('Professional Summary', level=2)
    para = add_para(doc,'','Results-driven Data Analyst with 6 years of experience in leveraging data to drive actionable insights and strategic decisions. Proficient in Data Fetching, Data cleansing, data preparation and data visualization using tools such as SQL, Excel, VBA, Python, Tableau and PowerBI. Adept at collaborating with cross-functional teams to identify key business challenges and implement data-driven solutions that optimize performance.')
   

    # Add Skills section with bar graph
    doc.add_heading('Skills', level=2)
    skills = ["Excel", "MySQL", "VBA", "Python", "Power BI", "Tableau"]
    ratings = [5, 4, 3, 3, 3, 3]
    max_rating = 5  # Maximum possible rating

    # Colors for the bars
    bar_colors = ['#009879'] * len(skills)  # Main bar color
    bg_colors = ['#b7e4c7'] * len(skills)  # Background bar color

    # Plot settings
    plt.figure(figsize=(4, 2))
    bar_height = 0.7  # Adjust height to reduce spacing (default is typically 0.8)
    for i, (skill, rating) in enumerate(zip(skills, ratings)):
        # Draw background bar
        plt.barh(y=i, width=max_rating, color=bg_colors[i], height=bar_height, edgecolor='none')
        # Draw rating bar
        plt.barh(y=i, width=rating, color=bar_colors[i], height=bar_height, edgecolor='none')
        # Add text for the rating and skill
        plt.text(rating - 0.9, i, f"{rating} / {max_rating}", va='center', ha='center', color='white', fontsize=9, weight='bold')
        plt.text(-0.99, i, skill, va='center', ha='left', color='black', fontsize=9)

    # Remove axes and labels for clean look
    plt.gca().invert_yaxis()  # Invert y-axis for descending order
    plt.axis('off')  # Hide the axes

    # Set background color
    #plt.gcf().patch.set_facecolor('#eaf4f4')  # Light blue background

    # Show or Save the chart as an image
    plt.tight_layout()
    #plt.show()
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
        chart_path = tmp_file.name
        plt.savefig(chart_path)
    plt.close()
    
    # Add the chart to the Word document
    doc.add_picture(chart_path, width=doc.sections[0].page_width * 0.4)  # Scale image to fit page
    
    # Add Professional Experience
    doc.add_heading('Professional Experience', level=2)
    doc.add_paragraph('Assistant Manager - MIS', style='Heading 3')
    add_company_details(doc,'','Standard Chartered Bank - GBS, Chennai, Tamil Nadu, 07/2023 – Present')
    add_bullet_point(doc, 'Optimized the Risk Team’s workflow ', 'by automating the consolidation of daily reports from a secured shared drive and extracting Reserve Accounts data using Power Query, significantly reducing manual effort and improving data accuracy.')
    add_bullet_point(doc, 'Enhanced the automated population arrival process ', 'for the Risk Team by designing solutions using Excel User Forms, VBA, and Power Query, leading to improved processing speed and efficiency.')
    add_bullet_point(doc, 'Developed a dynamic front-end tool in Excel ', 'for seamless source file uploads and efficient data retrieval from MS Access streamlining data management tasks.')
    add_bullet_point(doc, 'Automated data extraction and random sample selection processes ', 'boosting operational efficiency and ensuring greater accuracy in data processing.')
    add_bullet_point(doc, 'Designed and implemented automated reporting solutions ', 'for delivering daily and monthly reports, including Regular Reviews Trend Analysis, Monthly Burns, and Productivity, enabling more effective performance tracking and data-driven decision-making.')
  
    doc.add_page_break()
    
    doc.add_paragraph('Data Analyst', style='Heading 3')
    add_company_details(doc,'','Optum Health Care Business Services & Technology, Chennai, Tamil Nadu, 03/2021 – 07/2023')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, 'Managed extensive datasets ', 'by performing data extraction, manipulation, cleansing, quality assurance, and analysis to ensure data integrity and actionable insights.')
    add_bullet_point(doc, 'Contributed to project planning and process documentation ', 'encompassing methods, workflows, task scheduling, and resource management to facilitate seamless project execution.')
    add_bullet_point(doc, 'Collaborated with onshore teams, developers, and support staff ', 'to implement solutions for three healthcare clients, ensuring successful project delivery and client satisfaction.')
    add_bullet_point(doc, 'Executed QA test cases, ', 'monitored testing progress, and worked closely with developers to resolve issues prior to implementation, ensuring high-quality deliverables.')
    add_bullet_point(doc, 'Utilized SQL databases for data extraction, ', 'showcasing proficiency in querying and managing data to support various analytical tasks.')
    add_bullet_point(doc, 'Prepared and distributed daily, weekly, and monthly reports ', 'to clients and senior management, providing insights that supported informed decision-making.')

    doc.add_paragraph('Technical Process Specialist', style='Heading 3')
    add_company_details(doc,'','Infosys, Bangalore, Tamil Nadu, 09/2020 – 03/2021')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, 'Data Extraction and Analysis: ', 'Proficient in extracting, importing, and transforming data from diverse sources to derive actionable insights.')
    add_bullet_point(doc, 'Content Management Systems (CMS): ', 'Skilled in utilizing CMS platforms to export data tailored to specific requirements.')
    add_bullet_point(doc, 'Excel Reporting: ', 'Experienced in creating visual reports and dashboards using Excel to effectively present data.')
    add_bullet_point(doc, 'Access Database Management: ', 'Adept at creating, fetching, updating, and altering database tables in MySQL to meet project needs.')
    add_bullet_point(doc, 'Stakeholder Reporting: ', 'Experienced in developing and publishing reports for stakeholders, ensuring clear communication of findings.')

    doc.add_paragraph('Project Coordinator', style='Heading 3')
    add_company_details(doc,'','Origin Learning Solutions Pvt Ltd, Tamil Nadu, 10/2018 – 01/2020')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, 'Coordinate Project Schedules and Resources: ', 'Efficiently manage project timelines, allocate resources, and oversee requirements to ensure seamless project execution.')
    add_bullet_point(doc, 'Define Project Scope and Objectives: ', 'Collaborate with stakeholders to clearly outline project requirements, establish scope, and set achievable objectives.')
    add_bullet_point(doc, 'Assist in Preparing Project Proposals and Budgets: ', 'Support the development of comprehensive project proposals, including detailed timelines, schedules, and budget estimates.')
    add_bullet_point(doc, 'Monitor Project Progress and Address Issues: ', 'Continuously track project milestones, identify potential issues, and implement solutions to keep the project on course.')
    add_bullet_point(doc, 'Ensure Adherence to Project Deadlines: ', 'Proactively follow up with team members to confirm that project timelines are met and deliverables are completed as scheduled.')
    add_bullet_point(doc, 'Participate in Project Review Meetings and Risk Assessment: ', 'Engage actively in regular project evaluations, highlighting potential risks and contributing to mitigation strategies.')
    add_bullet_point(doc, 'Support Risk Identification and Contingency Planning: ', 'Collaborate with Delivery Managers and Heads to identify project risks and assist in developing effective contingency plans.')

    doc.add_page_break()

    doc.add_paragraph('Project Associate', style='Heading 3')
    add_company_details(doc,'','Emerson Automation Solutions Pvt Ltd, Tamil Nadu, 10/2018 – 01/2020')
    #add_bullet_point(doc, '', '')
    add_bullet_point(doc, 'Oversee PMO Activities: ', 'Manage Project Management Office (PMO) tasks from order entry through to delivery, ensuring alignment with organizational objectives.')
    add_bullet_point(doc, 'Participate in Project Initiation: ', 'Engage actively in kickoff meetings to comprehend project scope, delivery timelines, and critical requirements, facilitating clear communication among stakeholders. ')
    add_bullet_point(doc, 'Develop Delivery Schedules: ', 'Create comprehensive project delivery timelines and document submission schedules to ensure timely progress and adherence to deadlines.')
    add_bullet_point(doc, 'Manage Document Submissions: ', 'Submit technical documents to clients for approval, ensuring accuracy and compliance with project requirements.')
    add_bullet_point(doc, 'Address Client Feedback: ', 'Review and resolve customer comments promptly, maintaining positive client relationships and ensuring project specifications are met.')
    add_bullet_point(doc, 'Ensure Timely Documentation: ', 'Guarantee the on-time submission of documents, securing necessary approvals to keep the project on schedule.')
    add_bullet_point(doc, 'Coordinate Stakeholders: ', 'Collaborate with all stakeholders to ensure the timely completion of critical activities, including Bill of Materials (BOM) preparation, ordering, material receipt, assembly, inspection, and shipment.')
    add_bullet_point(doc, 'Monitor Project Progress: ', 'Actively participate in periodic project review meetings, highlighting potential risks and ensuring that project milestones are met. ')
    add_bullet_point(doc, 'Prepare Progress Reports: ', 'Compile and submit regular progress reports to customers, providing transparent updates on project status and any emerging issues.')

    # Add Education section
    doc.add_heading('Education', level=2)
    c = doc.add_paragraph('Bachelor of Engineering', style='Heading 3')
    add_para(doc,'','Engineering College, City, State, 07/2013 – 04/2014')
    # Add an indent to the bullet
    p_format = c.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    
    
    c = doc.add_paragraph('Higher Secondary School Certificate', style='Heading 3')
    add_para(doc,'','Higher Secondary School, City, State, 07/2012 – 04/2013')
    p_format = c.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)

    doc.add_heading('Additional Training or Certifications', level=2)
    c = doc.add_paragraph('MySQL for Data Analytics and Buisness Intelligence, 07/2020 - 08/2020, ', style='List Bullet')
    d = add_hyperlink(c,'Certificate Link' ,'https://www.udemy.com/')
    c = Custom_Style(c)

    c = doc.add_paragraph('Master Python by Coding 100 Practical Problems, 02/2024 - 03/2024, ', style='List Bullet')
    d = add_hyperlink(c,'Certificate Link' ,'https://www.udemy.com/')
    c = Custom_Style(c)

    # Add Languages (Optional)  
    doc.add_heading('Languages', level=2)
    c = doc.add_paragraph('Tamil (Proficient)', style='List Bullet')
    d = doc.add_paragraph('English (Proficient)', style='List Bullet')
    c = Custom_Style(c)
    d = Custom_Style(d)

    # Save the document as temp file to generate pdf
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        filename = tmp_file.name
        doc.save(filename)
    plt.close()
    
    #convert the docx temp file to pdf
    convert(filename, Resume_Name)
    
    #file_name = "Professional_Resume_Template.docx"
    #doc.save(file_name)
    #print(f"Resume template saved as {file_name}")

# Call the function to create the resume
create_resume()