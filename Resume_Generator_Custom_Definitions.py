import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from docx.shared import Pt, RGBColor, Cm
from docx2pdf import convert
from datetime import date
import tempfile
from docx.enum.style import WD_STYLE_TYPE


def add_horizontal_line(doc):
    last_para = doc.paragraphs[-1]
    p = last_para._element
    pPr = p.get_or_add_pPr()
    
    # Add or modify spacing element
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')  # Set space before to 0
    spacing.set(qn('w:after'), '0')   # Set space after to 0 (optional)
    pPr.append(spacing)
    
    # Add border
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')  # Adjust as necessary
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def document_colour(doc, doc_color):
    doc_color = doc_color.replace('#','')
    background = OxmlElement('w:background')
    background.set(qn('w:color'), doc_color)     # F3F9FB F0F8FA Define black background
    doc.element.insert(0, background)      # Insert it
    background_shp = OxmlElement('w:displayBackgroundShape') # Setting to use my background
    doc.settings.element.insert(0, background_shp)      # Apply setting
    
def header_custmizaion(header_para):
    # Customize the font properties for the header paragraph
    run = header_para.runs[0]  # Access the first run in the paragraph
    # Set font name, size, and color
    run.italic = True
    run.font.name = 'Calibri'  # Set desired font
    run.font.size = Pt(10)      # Set font size
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Set custom color (blue)
                
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    #new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    new_run.font.name = 'Calibri'
    new_run.font.size = Pt(10)
    new_run.font.color.rgb = RGBColor(79, 129, 189)
    new_run.font.underline = False

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

def picture_custoization(picture):
    # Set the space after the paragraph containing the picture
    picture_paragraph.paragraph_format.space_after = 0  # Reduce space after picture
    picture_paragraph.paragraph_format.space_before = 0  # Reduce space before picture


def Level1_customization(para):
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = para.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(13)
    para.bold = True
    para.font.color.rgb = RGBColor(97, 182, 205) #RGBColor(88, 50, 164) 


def Custom_Style(para):
    # Add an indent to the bullet
    p_format = para.paragraph_format
    p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
    
    para = para.runs[0]
    para.font.name = 'Calibri'
    para.font.size = Pt(11)
    para.font.color.rgb = RGBColor(0x00, 0x00, 0x00) 

    p_format.space_after = 0  # Minimize space after the paragraph before the line
    p_format.space_before = 0 # Minimize space before the paragraph before the line

def symbol_customization(symbol):
    symbol.font.name = "Segoe UI Emoji"
    symbol.font.size = Pt(10)
    symbol.font.color.rgb = RGBColor(97, 182, 205) #RGBColor(0x00, 0xB0, 0x50) #RGBColor(0, 176, 80)
    
def add_bullet_point(doc, bold_text, regular_text, font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
        para = doc.add_paragraph('', style='List Bullet')
        
        # Add an indent to the bullet
        p_format = para.paragraph_format
        p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
        p_format.first_line_indent = Cm(-0.5)  # Negative indent for hanging bullet
        
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        bold_font = bold_run.font
        bold_font.name = font_name
        bold_font.size = Pt(font_size)
        bold_font.color.rgb = RGBColor(*font_color)
        
        regular_run  = para.add_run(regular_text)
        regular_font = regular_run.font
        regular_font.name = font_name
        regular_font.size = Pt(font_size)
        regular_font.color.rgb = RGBColor(*font_color)
        
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format = para.paragraph_format
        paragraph_format.space_after = 0  # Minimize space after the paragraph before the line
        paragraph_format.space_before = 0 # Minimize space before the paragraph before the line

def add_para(doc, bold_text, regular_text, font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
    para = doc.add_paragraph('')
    
    # Add an indent to the bullet
    p_format = para.paragraph_format
    #p_format.left_indent = Cm(1.0)  # Indent bullet by 1 cm (adjust as needed)
    p_format.first_line_indent = Cm(1.0)  # Negative indent for hanging bullet
    
    bold_run = para.add_run(bold_text)
    bold_run.bold = True
    bold_font = bold_run.font
    bold_font.name = font_name
    bold_font.size = Pt(font_size)
    bold_font.color.rgb = RGBColor(*font_color)
    
    regular_run  = para.add_run(regular_text)
    regular_font = regular_run.font
    regular_font.name = font_name
    regular_font.size = Pt(font_size)
    regular_font.color.rgb = RGBColor(*font_color)

    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format = para.paragraph_format
    paragraph_format.space_after = 0  # Minimize space after the paragraph before the line
    paragraph_format.space_before = 0 # Minimize space before the paragraph before the line
    
def add_company_details(doc, bold_text, regular_text, date_from_to, company_info, font_name='Calibri', font_size=11, font_color=(0, 0, 0)):
    para = doc.add_paragraph('')
    
    bold_run = para.add_run(bold_text)
    bold_run.bold = True
    bold_font = bold_run.font
    bold_font.name = font_name
    bold_font.size = Pt(font_size)
    bold_font.color.rgb = RGBColor(*font_color)
    
    regular_run  = para.add_run(regular_text)
    regular_font = regular_run.font
    regular_font.name = font_name
    regular_font.size = Pt(font_size)
    regular_font.color.rgb = RGBColor(*font_color)
    
    
    date_from_to_run = para.add_run(date_from_to + "\n")
    date_from_to_run.bold = True
    date_from_to_run_font = date_from_to_run.font
    date_from_to_run_font.name = font_name
    date_from_to_run_font.size = Pt(font_size)
    date_from_to_run_font.color.rgb = RGBColor(*font_color)
    
    
    company_info_run = para.add_run(company_info)
    company_info_run.bold = False
    company_info_run.italic = True
    company_info_run_font = company_info_run.font
    company_info_run_font.name = font_name
    company_info_run_font.size = Pt(font_size)
    company_info_run_font.color.rgb = RGBColor(*font_color)

    #para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    